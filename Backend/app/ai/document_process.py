from pypdf import PdfReader
from docx import Document as DocxDocument

CHUNK_SIZE = 500       # approx. words per chunk
CHUNK_OVERLAP = 50     # words repeated between consecutive chunks, for context continuity


def _extract_text_from_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    pages_text = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages_text)


def _extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Word documents often contain tables too — extract those separately
    # so table content isn't silently skipped.
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def _chunk_text(text: str) -> list[str]:
    words = text.split()
    chunks = []
    start = 0

    while start < len(words):
        end = start + CHUNK_SIZE
        chunk_words = words[start:end]
        chunks.append(" ".join(chunk_words))
        start += CHUNK_SIZE - CHUNK_OVERLAP

    return chunks


def extract_and_chunk_text(file_path: str) -> list[str]:
    """
    Reads a document from disk and returns a list of text chunks,
    ready to be embedded and stored one by one.
    """
    lower_path = file_path.lower()

    if lower_path.endswith(".pdf"):
        raw_text = _extract_text_from_pdf(file_path)
    elif lower_path.endswith(".txt"):
        raw_text = _extract_text_from_txt(file_path)
    elif lower_path.endswith(".docx"):
        raw_text = _extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type for extraction: {file_path}")

    raw_text = raw_text.strip()
    if not raw_text:
        raise ValueError("No extractable text found in document")

    return _chunk_text(raw_text)